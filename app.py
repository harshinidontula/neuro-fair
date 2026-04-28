import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import docx
import json
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score
from fairlearn.metrics import selection_rate, demographic_parity_difference

# Page config
st.set_page_config(page_title="Neuro fair - Bias Platform", page_icon="⚖️", layout="wide")

# Custom UI styling
st.markdown("""
<style>
/* Main app background (white/light blue hint) */
.stApp {
    background-color: #f4f9ff;
}
/* Sidebar (dark professional blue) */
section[data-testid="stSidebar"] {
    background-color: #0e2a47;
}
section[data-testid="stSidebar"] * {
    color: #ffffff;
}
/* Main Headers */
h1, h2, h3 {
    color: #0e2a47;
    font-weight: 700;
}
/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, #0d47a1 0%, #1976d2 100%);
    color: white;
    border-radius: 8px;
    border: none;
    font-weight: 600;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
}
.stButton > button:hover {
    transform: scale(1.02);
    box-shadow: 0 4px 12px rgba(13, 71, 161, 0.4);
    color: white;
}
/* Metric Cards */
div[data-testid="stMetricValue"] {
    color: #1976d2;
    font-weight: bold;
}
/* Info/Success messages */
.stAlert {
    border-radius: 8px;
}
</style>
""", unsafe_allow_html=True)

# Helper config
@st.cache_data
def load_data(file):
    name = file.name.lower()
    try:
        if name.endswith('.csv') or name.endswith('.txt'):
            return pd.read_csv(file)
        elif name.endswith(('.xls', '.xlsx')):
            return pd.read_excel(file)
        elif name.endswith('.json'):
            return pd.read_json(file)
        elif name.endswith('.docx'):
            doc = docx.Document(file)
            if not doc.tables:
                st.error("No tables found in this Word document! Please provide a document with structured data tables.")
                return None
            table = doc.tables[0]
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(data) <= 1:
                st.error("Table in the document is too small or missing headers.")
                return None
            df = pd.DataFrame(data[1:], columns=data[0])
            for col in df.columns:
                df[col] = pd.to_numeric(df[col], errors='ignore')
            return df
        else:
            st.error(f"Unsupported file format: {name}. AI bias learning requires structured data (like CSV or Excel).")
            return None
    except Exception as e:
        st.error(f"Failed to read file: {str(e)}")
        return None


def train_eval_model(df_train, features_cols, target_col, sensitive_series):
    X = df_train[features_cols]
    y = df_train[target_col]
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

    clf = LogisticRegression(max_iter=1000)
    clf.fit(X_train, y_train)
    y_pred = clf.predict(X_test)
    acc = accuracy_score(y_test, y_pred)

    importance = clf.coef_[0]

    A_test = sensitive_series.loc[X_test.index]
    unique_groups = [g for g in A_test.unique() if str(g) != 'nan']
    sr_dict = {}
    for group in unique_groups:
        mask = (A_test == group)
        if mask.sum() > 0:
            sr_dict[str(group)] = selection_rate(y_test[mask], y_pred[mask])
        else:
            sr_dict[str(group)] = 0.0

    dpd = demographic_parity_difference(y_test, y_pred, sensitive_features=A_test)

    return acc, sr_dict, dpd, importance


# Sidebar Navigation
st.sidebar.title("⚖️ Neuro Fair (a FairAI Platform)")
menu = st.sidebar.radio("Navigation", [
    "1. Project Overview",
    "2. Data Upload",
    "3. Bias Detection",
    "4. Model Training & Fairness",
    "5. Bias Correction & Explainability",
    "6. Text Bias Detection (AI)"
])

# State variables
if 'df' not in st.session_state:
    st.session_state.df = None


# --- Page 1: Overview ---
if menu == "1. Project Overview":
    st.title("Welcome to FairAI")
    st.markdown("""
        **Algorithmic Bias and Fairness in Machine Learning**  
        FairAI is a comprehensive prototype for detecting, evaluating, and mitigating bias in tabular datasets.
        
        ### 🔍 How it Works:
        1. **Upload Dataset:** Provide a CSV file (like Loan Approvals or Hiring decisions).
        2. **Bias Detection:** Visualize imbalances in sensitive groups (e.g., Gender, Age).
        3. **Model & Fairness:** Train a predictive model and evaluate algorithmic fairness metrics (Selection Rate, Demographic Parity).
        4. **Bias Correction:** See how mitigating strategies (like removing sensitive proxies) adjust the model's fairness profile.
    """)
    st.image(
        "https://images.unsplash.com/photo-1551288049-bebda4e38f71?ixlib=rb-4.0.3&auto=format&fit=crop&w=1000&q=80",
        caption="Fairness in Decision Making",
        use_container_width=True
    )


# --- Page 2: Data Upload ---
elif menu == "2. Data Upload":
    st.title("Data Upload")
    uploaded_file = st.file_uploader(
        "Upload your dataset",
        type=["csv", "xlsx", "xls", "json", "txt", "docx", "pdf", "jpg", "jpeg", "png"],
        accept_multiple_files=False
    )

    st.info("💡 Note: While you can upload word docs or PDFs, the Bias Mitigation engine requires structured, tabular tables to detect mathematical unfairness. Use CSV or Excel for best results.")

    if uploaded_file is not None:
        name = uploaded_file.name.lower()
        is_image = name.endswith(('.jpg', '.jpeg', '.png'))
        is_pdf = name.endswith('.pdf')

        if is_image:
            st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)
            st.warning("⚠️ Image successfully uploaded. Note: The ML bias model requires structured tabular data. We've displayed the image, but it cannot currently be processed for bias metrics.")
            st.session_state.df = None

        elif is_pdf:
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text = ""
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"

                if not text.strip():
                    st.error("❌ We couldn't extract any text from this PDF. It appears to be a scanned image or protected document.")
                else:
                    st.session_state.unstructured_text = text
                    st.session_state.df = None

                    st.success("✅ Unstructured Text Document Detected!")
                    st.info("The PDF has been successfully read. Please navigate to **Page 6: Text Bias Detection (AI)** from the sidebar to analyze it for bias.")
                    with st.expander("Preview Extracted Text"):
                        st.text(text[:1000] + ("..." if len(text) > 1000 else ""))

            except ImportError:
                st.error("PyPDF2 is not installed. Please run `pip install PyPDF2`.")
            except Exception as e:
                st.error(f"Failed to read PDF: {str(e)}")

        else:
            df = load_data(uploaded_file)
            if df is not None:
                st.session_state.df = df
                st.success("File uploaded successfully!")
                st.dataframe(df.head())
                st.write(f"**Shape:** {df.shape[0]} rows, {df.shape[1]} columns")
    else:
        st.info("Please upload a file, or use the synthetic `dataset.csv` located in the `/data` folder.")


# --- Page 3: Bias Detection ---
elif menu == "3. Bias Detection":
    st.title("Bias Detection (EDA)")
    df = st.session_state.df

    if df is not None:
        sensitive_col = st.selectbox(
            "Select Sensitive Attribute",
            df.columns,
            index=list(df.columns).index("Gender") if "Gender" in df.columns else 0
        )
        target_col = st.selectbox(
            "Select Target Variable",
            df.columns,
            index=list(df.columns).index("Approved") if "Approved" in df.columns else len(df.columns) - 1
        )

        st.session_state.sensitive_col = sensitive_col
        st.session_state.target_col = target_col

        col1, col2 = st.columns(2)
        with col1:
            st.subheader(f"Distribution of {sensitive_col}")
            fig, ax = plt.subplots(figsize=(6, 5))
            df[sensitive_col].value_counts().plot.pie(autopct="%1.1f%%", ax=ax, cmap='Set3')
            ax.set_ylabel("")
            plt.tight_layout()
            st.pyplot(fig)

        with col2:
            st.subheader(f"Outcome ({target_col}) by {sensitive_col}")
            fig2, ax2 = plt.subplots(figsize=(8, 5))
            sns.countplot(data=df, x=sensitive_col, hue=target_col, palette='Set2', ax=ax2)
            ax2.tick_params(axis='x', rotation=45)
            plt.tight_layout()
            st.pyplot(fig2)

        st.markdown("### Observations")
        st.info("Visually inspect the charts to note any historical bias. If one group has a significantly lower positive outcome rate, the data is imbalanced and models trained on it will inherit this bias.")

    else:
        if st.session_state.get('unstructured_text'):
            st.warning("You have uploaded an unstructured text document (PDF). Please navigate to **Page 6: Text Bias Detection (AI)** from the sidebar to analyze it!")
        else:
            st.warning("Please upload a dataset in the Data Upload section.")


# --- Page 4: Model Training ---
elif menu == "4. Model Training & Fairness":
    st.title("Model Training & Fairness Evaluation")
    df = st.session_state.df

    if df is not None:
        sensitive_col = st.session_state.get('sensitive_col', df.columns[0])
        target_col = st.session_state.get('target_col', df.columns[-1])
        st.write(f"**Target:** `{target_col}` | **Sensitive Attribute:** `{sensitive_col}`")

        df_encoded = pd.get_dummies(df, drop_first=True)
        target_encoded_cols = [c for c in df_encoded.columns if target_col in c]

        if target_encoded_cols:
            true_target = target_encoded_cols[0]
            features = [c for c in df_encoded.columns if c != true_target]

            st.write("Training Logistic Regression Model...")
            acc, sr_dict, dpd, importance = train_eval_model(df_encoded, features, true_target, df[sensitive_col])

            st.session_state.baseline_metrics = {
                'acc': acc,
                'sr_dict': sr_dict,
                'dpd': dpd,
                'features': features,
                'importance': importance,
                'true_target': true_target
            }

            col1, col2, col3 = st.columns(3)
            col1.metric("Model Accuracy", f"{acc * 100:.1f}%")

            groups = list(sr_dict.keys())
            if len(groups) > 0:
                col2.metric(f"Selection Rate ({groups[0]})", f"{sr_dict[groups[0]] * 100:.1f}%")
            if len(groups) > 1:
                col3.metric(f"Selection Rate ({groups[1]})", f"{sr_dict[groups[1]] * 100:.1f}%")

            st.subheader("Demographic Parity Difference")
            st.write(f"The difference in selection rates is **{dpd * 100:.1f}%**.")
            st.warning("A high difference indicates potential bias.")
        else:
            st.error("Error processing target column.")

    else:
        if st.session_state.get('unstructured_text'):
            st.warning("You have uploaded an unstructured text document (PDF). Please navigate to **Page 6: Text Bias Detection (AI)** from the sidebar to analyze it!")
        else:
            st.warning("Please upload a dataset first.")


# --- Page 5: Correction & Explainability ---
elif menu == "5. Bias Correction & Explainability":
    st.title("Bias Mitigation & Explainable AI")
    df = st.session_state.df

    if df is not None and 'baseline_metrics' in st.session_state:
        st.markdown("### Fairness Intervention: Blinded Feature Selection")
        st.write("We will remove the sensitive attribute (and highly correlated proxies if chosen) and re-train the model.")

        sensitive_col = st.session_state.get('sensitive_col')
        target_col = st.session_state.get('target_col')
        df_encoded = pd.get_dummies(df, drop_first=True)
        true_target = st.session_state.baseline_metrics['true_target']

        sensitive_feat_cols = [c for c in df_encoded.columns if sensitive_col in c]
        mitigated_features = [c for c in df_encoded.columns if c not in sensitive_feat_cols and c != true_target]

        if st.button("Apply Mitigation & Re-train"):
            m_acc, m_sr_dict, m_dpd, m_imp = train_eval_model(df_encoded, mitigated_features, true_target, df[sensitive_col])

            st.subheader("Before vs After Mitigation")

            groups = list(m_sr_dict.keys())
            metrics_list = ["Accuracy"]
            before_list = [f"{st.session_state.baseline_metrics['acc'] * 100:.1f}%"]
            after_list = [f"{m_acc * 100:.1f}%"]

            for g in groups[:2]:
                metrics_list.append(f"Selection Rate ({g})")
                b_val = st.session_state.baseline_metrics['sr_dict'].get(g, 0)
                before_list.append(f"{b_val * 100:.1f}%")
                after_list.append(f"{m_sr_dict[g] * 100:.1f}%")

            metrics_list.append("Demographic Parity Diff")
            before_list.append(f"{st.session_state.baseline_metrics['dpd'] * 100:.1f}%")
            after_list.append(f"{m_dpd * 100:.1f}%")

            comp_df = pd.DataFrame({
                "Metric": metrics_list,
                "Before": before_list,
                "After": after_list
            })
            st.table(comp_df.set_index("Metric"))

            st.success("Notice how the Selection Rates have ideally moved closer together, dropping Demographic Parity Difference.")

            st.markdown("---")
            st.subheader("Explainable AI: Feature Importance")
            st.write("Which features are driving the mitigated model's decisions?")

            fig, ax = plt.subplots(figsize=(8, 4))
            importance_df = pd.DataFrame({
                'Feature': mitigated_features,
                'Importance': m_imp
            }).sort_values(by='Importance', ascending=False)
            sns.barplot(data=importance_df, x='Importance', y='Feature', palette='viridis', ax=ax)
            st.pyplot(fig)

    else:
        if df is None and 'unstructured_text' in st.session_state:
            st.warning("You have uploaded an unstructured text document (PDF). Please navigate to **Page 6: Text Bias Detection (AI)** from the sidebar!")
        else:
            st.warning("Please complete Model Training (Step 4) first.")


# --- Page 6: Text Bias Detection ---
elif menu == "6. Text Bias Detection (AI)":
    st.title("Unstructured Text Bias Detection")
    st.markdown("Use our advanced AI to analyze text (e.g., job descriptions, emails, performance reviews) for potential bias.")

    st.subheader("1. Input Document")
    default_text = st.session_state.get('unstructured_text', "")
    user_text = st.text_area(
        "Enter text to analyze:",
        value=default_text,
        height=200,
        placeholder="Paste your document or paragraph here...",
        label_visibility="collapsed"
    )

    analyze_btn = st.button("Analyze Document", type="primary")

    st.markdown("---")

    if analyze_btn:
        st.subheader("2. AI Analysis Dashboard")

        # Load API key from secrets
        try:
            api_key = st.secrets["GEMINI_API_KEY"]
        except Exception:
            st.error("🚨 **Backend Configuration Error:** The `GEMINI_API_KEY` is not set in `.streamlit/secrets.toml`.")
            st.stop()

        if not api_key or api_key == "your-gemini-api-key-here":
            st.error("🚨 **Backend Configuration Error:** Please place your real Gemini API key in `.streamlit/secrets.toml`.")
            st.stop()

        if not user_text.strip():
            st.warning("Please enter some text to analyze.")
        else:
            with st.spinner("Backend AI is processing the document..."):
                try:
                    import google.generativeai as genai
                    genai.configure(api_key=api_key)

                    # ✅ Fixed: valid Gemini model name
                    model = genai.GenerativeModel('gemini-1.5-flash')

                    prompt = f"""You are an advanced Bias Detection AI.

Your task is to analyze unstructured text (paragraphs, documents, or mixed content) and identify any form of bias.

Instructions:
1. Break the input text into individual sentences.
2. Analyze each sentence independently while preserving context.
3. Detect the following types of bias:
   - Gender Bias
   - Racial Bias
   - Cultural Bias
   - Religious Bias
   - Socioeconomic Bias
   - Age Bias
   - Stereotyping or Discrimination

4. For each sentence, return:
   - sentence
   - bias_detected (Yes/No)
   - bias_type (if any, else "None")
   - confidence_score (0 to 1)
   - explanation (short and clear reason)

5. If no bias is found, clearly state "No bias detected".

6. Output MUST be in strict JSON format like this:

[
  {{
    "sentence": "...",
    "bias_detected": "Yes/No",
    "bias_type": "...",
    "confidence_score": 0.85,
    "explanation": "..."
  }}
]

7. Do NOT skip any sentence.
8. Do NOT add extra text outside JSON.

Now analyze the following text:

{user_text}"""

                    response = model.generate_content(prompt)
                    output_text = response.text.strip()

                    # Clean up JSON formatting if wrapped in markdown code blocks
                    if output_text.startswith("```json"):
                        output_text = output_text[7:]
                    if output_text.startswith("```"):
                        output_text = output_text[3:]
                    if output_text.endswith("```"):
                        output_text = output_text[:-3]

                    results = json.loads(output_text.strip())

                    if not results:
                        st.info("No sentences found in the input text.")
                    else:
                        st.success("Analysis Complete!")
                        st.markdown("---")

                        for idx, item in enumerate(results):
                            if str(item.get("bias_detected", "No")).lower().strip() in ["yes", "true", "1"]:
                                st.error(f"**Sentence {idx + 1}:** {item.get('sentence')}")
                                st.write(f"⚠️ **Bias Type:** {item.get('bias_type')} | **Confidence:** {item.get('confidence_score')}")
                                st.write(f"📝 **Explanation:** {item.get('explanation')}")
                                st.markdown("---")
                            else:
                                st.success(f"**Sentence {idx + 1}:** {item.get('sentence')}")
                                st.write("✅ **No Bias Detected**")
                                st.write(f"📝 **Explanation:** {item.get('explanation')}")
                                st.markdown("---")

                except ImportError:
                    st.error("The `google-generativeai` library is not installed. Please run `pip install google-generativeai`.")
                except json.JSONDecodeError:
                    st.error("Failed to parse the AI output. The model returned invalid JSON.")
                    if 'output_text' in locals():
                        with st.expander("Show raw output"):
                            st.write(output_text)
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
